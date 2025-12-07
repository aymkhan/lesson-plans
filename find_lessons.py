# Aymen Khan
# Dec 2025

# This program will find the Lessons and Activities for ASL1, ASL2, and ASL3 courses

import os
import csv
import pandas as pd
from docx import Document

def start():
	print("Hello, Gianna!")

ASL1_LESSON_KEYS = []
ASL1_LESSON_VALUES = {}
ASL1_LESSONS = {}

ASL1_file = os.getcwd() + "/database/lessons/ASL1_Lessons.csv"

with open(ASL1_file,'r') as asl1:
	data = pd.DataFrame(asl1)
	print(data)
#	for column in asl1:
#		if "Lesson 2" in column:
#			print(column)


'''
Lines below are responsible for working with DOCX files

# Set the paths for LESSON .docx files
ASL1_docx = os.getcwd() + "/database/lessons/ASL1_lessons.docx"
print(f"os.path: ", ASL1_docx) 

# Open the .docx file to read
ASL1 = Document(ASL1_docx)
p = ASL1.paragraphs

lesson = "Lesson"
counter = 0

for table in ASL1.tables:
	for row in table.rows:
		for cell in row.cells:
			if lesson in cell.text:
				counter +=1
				lesson = "Lesson " + str(counter)
				print("Found lesson ", counter)
				ASL1_LESSON_KEYS.append(lesson)


for line in p:
	print(line.text) #prints each line
	if lesson in line.text:
		counter +=1
		lesson = "Lesson " + str(counter)
		print("Found lesson ", counter)
		ASL1_LESSON_KEYS.append(lesson)


print(f"Lesson keys: ", ASL1_LESSON_KEYS)
'''
		

